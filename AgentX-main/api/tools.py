import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

logger = logging.getLogger(__name__)

class BudgetPlanner:
    """Generate personalized budget reports in DOCX/PDF format"""
    
    def __init__(self, output_dir: str = "./data/reports"):
        self.output_dir = output_dir
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        logger.info(f"✅ BudgetPlanner initialized (output: {output_dir})")
    
    def generate(self, user_id: str, income: float, 
                expenses: Dict[str, float], output_format: str = "docx") -> str:
        """Generate budget report"""
        logger.info(f"📊 Generating {output_format} budget report...")
        
        if output_format == "docx":
            return self._generate_docx(user_id, income, expenses)
        else:
            return self._generate_pdf(user_id, income, expenses)
    
    def _generate_docx(self, user_id: str, income: float, 
                      expenses: Dict[str, float]) -> str:
        """Generate DOCX report"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Your Monthly Budget Plan', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"User ID: {user_id}")
        
        # Income Section
        doc.add_heading('Income Summary', level=1)
        doc.add_paragraph(f"Monthly Income: Rs {income:,.2f}")
        
        # Expenses Section
        doc.add_heading('Expense Breakdown', level=1)
        
        total_expenses = sum(expenses.values())
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Amount (Rs)'
        hdr_cells[2].text = 'Percentage'
        
        # Data rows
        for category, amount in expenses.items():
            row_cells = table.add_row().cells
            percentage = (amount / income) * 100 if income > 0 else 0
            row_cells[0].text = category.capitalize()
            row_cells[1].text = f"{amount:,.2f}"
            row_cells[2].text = f"{percentage:.1f}%"
        
        # Total row
        row_cells = table.add_row().cells
        row_cells[0].text = "TOTAL"
        row_cells[1].text = f"{total_expenses:,.2f}"
        percentage = (total_expenses / income) * 100 if income > 0 else 0
        row_cells[2].text = f"{percentage:.1f}%"
        
        # Summary
        doc.add_heading('Budget Summary', level=1)
        surplus = income - total_expenses
        doc.add_paragraph(f"Monthly Income: Rs {income:,.2f}")
        doc.add_paragraph(f"Total Expenses: Rs {total_expenses:,.2f}")
        doc.add_paragraph(f"Monthly Surplus: Rs {surplus:,.2f}")
        
        # 50/30/20 Rule
        doc.add_heading('50/30/20 Budgeting Rule', level=1)
        doc.add_paragraph(
            "A recommended budget allocation method that divides your income into three categories:",
            style='List Bullet'
        )
        
        needs = income * 0.5
        wants = income * 0.3
        savings = income * 0.2
        
        rule_table = doc.add_table(rows=4, cols=2)
        rule_table.style = 'Light Grid Accent 1'
        
        rule_table.rows[0].cells[0].text = "Category"
        rule_table.rows[0].cells[1].text = "Amount (Rs)"
        
        rule_table.rows[1].cells[0].text = "Needs (50%)"
        rule_table.rows[1].cells[1].text = f"{needs:,.2f}"
        
        rule_table.rows[2].cells[0].text = "Wants (30%)"
        rule_table.rows[2].cells[1].text = f"{wants:,.2f}"
        
        rule_table.rows[3].cells[0].text = "Savings (20%)"
        rule_table.rows[3].cells[1].text = f"{savings:,.2f}"
        
        # Emergency Fund
        doc.add_heading('Emergency Fund Target', level=1)
        emergency_fund = total_expenses * 6
        monthly_target = emergency_fund / 12
        doc.add_paragraph(f"Recommended Emergency Fund: Rs {emergency_fund:,.2f}")
        doc.add_paragraph(f"(Covers {int(emergency_fund/total_expenses)} months of expenses)")
        doc.add_paragraph(f"Monthly Savings Target: Rs {monthly_target:,.2f}")
        doc.add_paragraph(f"Timeline to Achieve (in months): 12")
        
        # Tips
        doc.add_heading('Budget Tips', level=1)
        tips = [
            "Track all expenses for a week to verify these numbers",
            "Cut down on wants to increase savings",
            "Automate savings by setting up transfers on payday",
            "Review and adjust your budget quarterly",
            "Gradually increase your savings rate"
        ]
        for tip in tips:
            doc.add_paragraph(tip, style='List Bullet')
        
        # Save file
        filename = f"budget_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        logger.info(f"✅ DOCX saved: {file_path}")
        return file_path
    
    def _generate_pdf(self, user_id: str, income: float, 
                     expenses: Dict[str, float]) -> str:
        """Generate PDF report with charts"""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, PageBreak, Image, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        logger.info("📊 Generating PDF with visualizations...")
        
        # Create charts
        total_expenses = sum(expenses.values())
        
        fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(12, 10))
        fig.suptitle('Financial Overview', fontsize=16, fontweight='bold')
        
        # Chart 1: Expense Pie Chart
        colors_list = plt.cm.Set3(range(len(expenses)))
        ax1.pie(expenses.values(), labels=expenses.keys(), autopct='%1.1f%%', colors=colors_list)
        ax1.set_title('Expense Distribution')
        
        # Chart 2: Income vs Expenses
        categories = ['Expenses', 'Surplus']
        amounts = [total_expenses, income - total_expenses]
        colors_bar = ['#ff9999' if x < 0 else '#90EE90' for x in amounts]
        ax2.bar(categories, amounts, color=colors_bar)
        ax2.set_ylabel('Amount (Rs)')
        ax2.set_title('Income vs Expenses')
        ax2.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
        
        # Chart 3: 50/30/20 Rule
        rule_categories = ['Needs\n(50%)', 'Wants\n(30%)', 'Savings\n(20%)']
        rule_amounts = [income * 0.5, income * 0.3, income * 0.2]
        rule_colors = ['#87CEEB', '#FFD700', '#90EE90']
        ax3.bar(rule_categories, rule_amounts, color=rule_colors)
        ax3.set_ylabel('Amount (Rs)')
        ax3.set_title('50/30/20 Budget Rule')
        
        # Chart 4: Emergency Fund Progress
        months = list(range(0, 13))
        emergency_target = total_expenses * 6
        monthly_savings = emergency_target / 12
        savings_progress = [monthly_savings * m for m in months]
        
        ax4.plot(months, savings_progress, marker='o', linewidth=2, color='#FF6B6B')
        ax4.axhline(y=emergency_target, color='green', linestyle='--', linewidth=2, label='Target')
        ax4.fill_between(months, 0, savings_progress, alpha=0.3)
        ax4.set_xlabel('Months')
        ax4.set_ylabel('Amount Saved (Rs)')
        ax4.set_title('Emergency Fund Building Timeline')
        ax4.legend()
        ax4.grid(True, alpha=0.3)
        
        plt.tight_layout()
        chart_path = os.path.join(self.output_dir, f"chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        # Create PDF
        filename = f"budget_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(self.output_dir, filename)
        
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=1
        )
        story.append(Paragraph("Your Personal Budget Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Summary
        story.append(Paragraph("Budget Summary", styles['Heading2']))
        summary_data = [
            ['Metric', 'Amount'],
            ['Monthly Income', f'Rs {income:,.0f}'],
            ['Total Expenses', f'Rs {total_expenses:,.0f}'],
            ['Monthly Surplus', f'Rs {income - total_expenses:,.0f}'],
            ['Emergency Fund (6 months)', f'Rs {total_expenses * 6:,.0f}'],
        ]
        summary_table = Table(summary_data)
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Charts
        story.append(Paragraph("Financial Visualizations", styles['Heading2']))
        img = Image(chart_path, width=7*inch, height=5.25*inch)
        story.append(img)
        
        # Build PDF
        doc.build(story)
        
        # Cleanup
        os.remove(chart_path)
        
        logger.info(f"✅ PDF saved: {file_path}")
        return file_path
